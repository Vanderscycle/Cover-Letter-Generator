from docx import Document
from docx.shared import Inches, Cm
# from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt 
import datetime
import json, re, os
from pprint import pprint
# custom made python script
import helper
# dotenv to handle the .env file
from dotenv import load_dotenv
# pathlib to handle the 
from pathlib import Path
############## CONFIG
# detail file goes here
info = "sample.json"
save_dir = "coverLetters/"

BASE_DIR = Path(__file__).resolve().parent
savePath = os.path.join(BASE_DIR,save_dir)

# loading the json file data
data = json.load(open(info))

# loading all the private information
load_dotenv()
name =os.getenv('NAME')
address = os.getenv('ADDRESS')
phone = os.getenv('PHONE')
mail = os.getenv('MAIL')
website = os.getenv('WEBSITE')
github = os.getenv('GITHUB')

# updating the data to the json object from the .env file
data.update({
    'name':name,
    'address':address,
    'phone':phone,
    'mail':mail,
    'website':website,
    'github':github
}) 
  
# the result is a JSON string: 
# pprint(json.dumps(data))  

data["company_name"] = helper.askInput("Enter name of the company:")
company_name=data["company_name"]
data["position"] = helper.askInput("Enter the position name: ")
position=data["position"]

open_para=data["open_para"]
first_coop=data["first_coop"]
second_coop=data["second_coop"]
third_coop=data["third_coop"]
activities=data["activities"]
closing=data["closing"]

document = Document()

document.add_heading(name, 2)

margin = 2
#changing the page margins
sections = document.sections
for section in sections:
    section.top_margin = Cm(margin)
    section.bottom_margin = Cm(margin)
    section.left_margin = Cm(margin)
    section.right_margin = Cm(margin)


contact_info = document.add_paragraph()



#Contact information
address_obj = contact_info.add_run(address + " , ")
phone_obj = contact_info.add_run(phone + " , ")
mail_obj = contact_info.add_run(mail)
website_obj = contact_info.add_run("\n" + "Website:" + website)
github_obj = contact_info.add_run("\n" + "Github:" + github)

document.add_heading("_" * 90, 6)

# datetime
today = datetime.datetime.today()
date_obj =  document.add_paragraph().add_run(today.strftime('%d, %b %Y'))


re_obj = document.add_paragraph()
re_obj.add_run("Re: " + position)

dear_obj = document.add_paragraph()
dear_obj.add_run("Dear Hiring Manager,")

#Open paragraph
openpara_obj = document.add_paragraph()
openpara_obj = helper.format_alignment(openpara_obj)
open_para_choice = helper.askForChoice(open_para, "Choose open paragraph")
openpara_obj.add_run(helper.format_fill_in_info(open_para_choice, data))

#Third coop
third_coop_obj = document.add_paragraph()
third_coop_obj = helper.format_alignment(third_coop_obj)
third_coop_obj.add_run(third_coop)

#Second coop
second_coop_obj = document.add_paragraph()
second_coop_obj.add_run(second_coop)
second_coop_obj = helper.format_alignment(second_coop_obj)

#First coop
first_coop_obj = document.add_paragraph()
first_coop_obj = helper.format_alignment(first_coop_obj)
first_coop_obj.add_run(first_coop)

##############ACTIVITIES
activities_obj = document.add_paragraph()
activities_obj = helper.format_alignment(activities_obj)
activities = helper.askForChoices(activities, "Choose extracurricular activities")
activity_string = "Beside coop experience, {}. Additionally, {}. {}. Last but not least, {}."
activiy_para = activity_string.format(*activities)

activities_obj.add_run(activiy_para)

##############FINAL PARAGRAPH
final_obj = document.add_paragraph()
final_obj = helper.format_alignment(final_obj)
final_obj.add_run(helper.format_fill_in_info(closing, data))
   

##############CLOSING
closing_obj= document.add_paragraph()
closing_obj = helper.format_alignment(closing_obj, 0)
closing_obj.add_run("Sincerely,")

##############SIGNATURE
name_obj= document.add_paragraph()
name_obj = helper.format_alignment(name_obj, 0)
name_obj.add_run(name)

# save_file_name = company_name.strip(" ") + "-" + position.strip(" ") + today.strftime('%d, %b %Y').strip(" ")  +  ".docx"

save_file_name = helper.sanitize_name([company_name, position, today.strftime('%d, %b %Y') ])

save_to_path = savePath + save_file_name

print("Save to:", save_to_path)

document.save(save_to_path)

# os.system("abiword --to=pdf " + save_to_path) # Convert to pdf using abiword
# os.system("rm -rf " + save_to_path) # Delete the doc file

# clearing all the info from the .env 
data.update({
    'name':'',
    'address':'',
    'phone':'',
    'mail':'',
    'website':'',
    'github':''
}) 

with open(info, 'w') as outfile:
    json.dump(data, outfile, indent=4)
